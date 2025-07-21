"""
Document ingestion and processing for the One Ring platform.

This module provides functionality to ingest and process various document formats
(PDF, EPUB, DOCX, TXT) and extract text content for further processing.
"""

import io
import logging
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import magic
import pandas as pd
from bs4 import BeautifulSoup
from loguru import logger
from pydantic import BaseModel, Field

from one_ring.core.config import settings
from one_ring.utils.cleanup import register_cleanup_handler

# Try to import optional dependencies
try:
    import PyPDF2
    import docx
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    HAS_DOC_DEPS = True
except ImportError:
    HAS_DOC_DEPS = False


class DocumentMetadata(BaseModel):
    """Metadata for a document."""
    title: Optional[str] = None
    author: Optional[str] = None
    language: Optional[str] = "en"
    source: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    extra: Dict[str, Any] = Field(default_factory=dict)


class DocumentChunk(BaseModel):
    """A chunk of text from a document with metadata."""
    text: str
    metadata: DocumentMetadata
    page_number: Optional[int] = None
    chunk_number: int = 0
    chunk_size: int = 0
    chunk_overlap: int = 0
    
    class Config:
        """Pydantic config."""
        json_encoders = {
            Path: lambda v: str(v) if v else None,
        }


class DocumentProcessor:
    """Process documents and extract text content."""
    
    def __init__(self, chunk_size: int = 2000, chunk_overlap: int = 200):
        """Initialize the document processor.
        
        Args:
            chunk_size: Maximum size of each text chunk in characters.
            chunk_overlap: Number of characters to overlap between chunks.
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = [
            'application/pdf',
            'application/epub+zip',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain',
            'text/markdown',
            'text/csv',
            'text/tab-separated-values',
            'application/json',
        ]
        
        # Register cleanup handler
        register_cleanup_handler(self.cleanup)
        
        logger.info(f"Document processor initialized with chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
    
    def detect_file_type(self, file_path: Union[str, Path]) -> Optional[str]:
        """Detect the MIME type of a file.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            The MIME type of the file, or None if it cannot be determined.
        """
        try:
            # Use python-magic to detect the file type
            mime = magic.Magic(mime=True)
            return mime.from_file(str(file_path))
        except Exception as e:
            logger.warning(f"Failed to detect file type for {file_path}: {e}")
            return None
    
    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """Check if a file is supported by the document processor.
        
        Args:
            file_path: Path to the file.
            
        Returns:
            True if the file is supported, False otherwise.
        """
        file_type = self.detect_file_type(file_path)
        return file_type in self.supported_formats
    
    def process_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None,
    ) -> List[DocumentChunk]:
        """Process a file and extract text chunks.
        
        Args:
            file_path: Path to the file to process.
            metadata: Additional metadata to include with the document.
            
        Returns:
            A list of DocumentChunk objects.
            
        Raises:
            ValueError: If the file type is not supported.
            IOError: If the file cannot be read.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file type
        file_type = self.detect_file_type(file_path)
        if file_type not in self.supported_formats:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Initialize metadata
        doc_metadata = DocumentMetadata(
            source=str(file_path.name),
            **({} if metadata is None else metadata)
        )
        
        # Process based on file type
        if file_type == 'application/pdf':
            return self._process_pdf(file_path, doc_metadata)
        elif file_type == 'application/epub+zip':
            return self._process_epub(file_path, doc_metadata)
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self._process_docx(file_path, doc_metadata)
        elif file_type in ('text/plain', 'text/markdown'):
            return self._process_text(file_path, doc_metadata)
        elif file_type in ('text/csv', 'text/tab-separated-values'):
            return self._process_csv(file_path, doc_metadata)
        elif file_type == 'application/json':
            return self._process_json(file_path, doc_metadata)
        else:
            raise ValueError(f"Unhandled file type: {file_type}")
    
    def _process_pdf(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process a PDF file."""
        if not HAS_DOC_DEPS:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install pypdf2")
        
        chunks = []
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Update metadata
                metadata.page_count = len(pdf_reader.pages)
                if not metadata.title and hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                    metadata.title = pdf_reader.metadata.get('/Title', '')
                    metadata.author = pdf_reader.metadata.get('/Author', '')
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            page_chunks = self._chunk_text(
                                text=page_text,
                                metadata=metadata,
                                page_number=page_num,
                            )
                            chunks.extend(page_chunks)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num} of {file_path}: {e}")
                
                return chunks
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    def _process_epub(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process an EPUB file."""
        if not HAS_DOC_DEPS:
            raise ImportError("EbookLib is required for EPUB processing. Install with: pip install EbookLib")
        
        chunks = []
        
        try:
            book = epub.read_epub(str(file_path))
            
            # Update metadata
            if not metadata.title and book.get_metadata('DC', 'title'):
                metadata.title = book.get_metadata('DC', 'title')[0][0]
            if not metadata.author and book.get_metadata('DC', 'creator'):
                metadata.author = book.get_metadata('DC', 'creator')[0][0]
            if not metadata.language and book.get_metadata('DC', 'language'):
                metadata.language = book.get_metadata('DC', 'language')[0][0]
            
            # Extract text from each document in the EPUB
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text = soup.get_text()
                    
                    if text.strip():
                        doc_chunks = self._chunk_text(
                            text=text,
                            metadata=metadata,
                        )
                        chunks.extend(doc_chunks)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing EPUB {file_path}: {e}")
            raise
    
    def _process_docx(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process a DOCX file."""
        if not HAS_DOC_DEPS:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text += " | ".join(row_text) + "\n"
            
            return self._chunk_text(text=text, metadata=metadata)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    def _process_text(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process a plain text or markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return self._chunk_text(text=text, metadata=metadata)
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise
    
    def _process_csv(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process a CSV or TSV file."""
        try:
            # Read the CSV file
            df = pd.read_csv(file_path)
            
            # Convert to markdown table
            text = df.to_markdown(index=False)
            
            return self._chunk_text(text=text, metadata=metadata)
            
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {e}")
            raise
    
    def _process_json(self, file_path: Path, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process a JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to formatted JSON string
            text = json.dumps(data, indent=2, ensure_ascii=False)
            
            return self._chunk_text(text=text, metadata=metadata)
            
        except Exception as e:
            logger.error(f"Error processing JSON file {file_path}: {e}")
            raise
    
    def _chunk_text(
        self,
        text: str,
        metadata: DocumentMetadata,
        page_number: Optional[int] = None,
    ) -> List[DocumentChunk]:
        """Split text into chunks with overlap.
        
        Args:
            text: The text to split into chunks.
            metadata: The document metadata.
            page_number: Optional page number for the text.
            
        Returns:
            A list of DocumentChunk objects.
        """
        if not text.strip():
            return []
        
        # Clean up the text
        text = self._clean_text(text)
        
        # Split into chunks
        chunks = []
        start = 0
        chunk_num = 0
        
        while start < len(text):
            # Calculate end position
            end = min(start + self.chunk_size, len(text))
            
            # If we're not at the end, try to find a good breaking point
            if end < len(text):
                # Look for the last newline or space in the chunk
                break_pos = text.rfind('\n', start, end)
                if break_pos == -1 or break_pos < start + self.chunk_size // 2:
                    # No good newline found, try to break at a space
                    break_pos = text.rfind(' ', start, end)
                
                if break_pos > start + self.chunk_size // 2:
                    end = break_pos + 1
            
            # Extract the chunk
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk = DocumentChunk(
                    text=chunk_text,
                    metadata=metadata.copy(update={
                        "page_number": page_number,
                        "chunk_number": chunk_num,
                        "chunk_size": len(chunk_text),
                        "chunk_overlap": self.chunk_overlap,
                    }),
                    page_number=page_number,
                    chunk_number=chunk_num,
                    chunk_size=len(chunk_text),
                    chunk_overlap=self.chunk_overlap,
                )
                chunks.append(chunk)
                chunk_num += 1
            
            # Move to the next chunk with overlap
            start = end - self.chunk_overlap
            if start <= 0:
                start = end
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean up text by removing extra whitespace and normalizing."""
        # Replace various whitespace characters with a single space
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
        
        # Normalize unicode
        text = text.encode('utf-8', 'ignore').decode('utf-8')
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def cleanup(self) -> None:
        """Clean up any resources used by the document processor."""
        # Nothing to clean up in this implementation
        pass


def get_document_processor() -> DocumentProcessor:
    """Get a document processor instance."""
    return DocumentProcessor()


# Create a global document processor instance
document_processor = DocumentProcessor()

# Register cleanup handler
import atexit
atexit.register(document_processor.cleanup)

__all__ = ["DocumentProcessor", "DocumentChunk", "DocumentMetadata", "document_processor"]
